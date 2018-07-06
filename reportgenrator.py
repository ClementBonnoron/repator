import re
import json
from collections import OrderedDict
from docx import Document
from docx.shared import Cm

class Generator:
    def __escape_str(s):
        return s.replace("\a", "\\a").replace("\b", "\\b").replace("\f", "\\f").replace("\r", "\\r").replace("\t", "\\t").replace("\v", "\\v")

    def generate_report(json_content):
        if type(json_content) is str:
            return json_content.replace("%", "\%")

        if "name" not in json_content:
            json_content["name"] = ""
        if "content" not in json_content:
            json_content["content"] = ""

        file_content = None
        with open("templates/" + json_content["type"] + ".tex") as f:
            file_content = f.read()
        final_content = file_content
        final_content = re.sub("##.*NAME##", json_content["name"],final_content)

        content = ""
        if type(json_content["content"]) is list:
            for e in json_content["content"]:
                content += Generator.generate_report(e)
        elif type(json_content["content"]) is dict:
            content += Generator.generate_report(json_content["content"])
        elif type(json_content["content"]) is str:
            content += Generator.generate_report(json_content["content"])

        final_content = re.sub("##.*CONTENT##", content,final_content)
        return Generator.__escape_str(final_content)

    def __sub_dict(d, text):
        for k in d:
            if type(d[k]) is str:
                text = re.sub("##"+k+"##", d[k], text)
        return text

    def __do_fill(d, content):
        if type(d) is str:
            d = Generator.__sub_dict(content, d)
            return d
        if "content" not in d:
            return d
        if d["type"] == "unordered_list" or d["type"] == "ordered_list":
            l = []
            for e in content[d["filer"]]:
                template = dict(d["content"][0])
                template["content"] = Generator.__do_fill(template["content"],e)
                l.append(template)
            d["content"] = l
            return d
        if "filer" in d:
            content = content[d["filer"]]
        if type(d["content"]) is list:
            l = []
            for e in d["content"]:
                l.append(Generator.__do_fill(e, content))
            d["content"] = l
        elif type(d["content"]) is str:
            d["content"] = Generator.__sub_dict(content, d["content"])
        else:
                d["content"] = Generator.__do_fill(d["content"], content)
        return d


    def generate_json(json_content):
        structure = None
        with open("content/structure.json", "r") as f:
            structure = f.read()
            structure = json.loads(structure)

        result_json = {}
        result_json["type"] = "report"
        result_json["content"] = []
        for e in structure:
            with open("content/" + e + ".json", "r") as f:
                file_content = f.read()
            json_file_content = json.loads(file_content)
            res = Generator.__do_fill(json_file_content, json_content)
            result_json["content"].append(res)
        return result_json

    def __cut_before(string, match):
        return string[string.find(match)+len(match):]

    def __cut_after(string, match):
        return string[0:string.find(match)]

    def __get_body(xml_content):
        out = Generator.__cut_before(xml_content, "<w:body>")
        out = Generator.__cut_after(out, "</w:body>")
        return out

    def generate_docx(document, json):

        if isinstance(json, str):
            document.text = json

        if "type" in json:
            if json["type"] == "table":
                table = document.add_table(json["row"], json["col"], json["style"])
                for row in range(0,json["row"]):
                    for col in range(0,json["col"]):
                        Generator.generate_docx(table.cell(row, col),
                                                json["content"][row][col])
                if "width" in json:
                    col = 0
                    for width in json["width"]:
                        table.columns[col].width = Cm(width)
                        col += 1

            if json["type"] == "document":
                newDoc = Document(json["path"])

                for paragraph in newDoc.paragraphs:
                    text = paragraph.text
                    style = paragraph.style.name
                    p = document.add_paragraph(text, style)
                    p.paragraph_format = paragraph.paragraph_format

        if "name" in json:
            document.add_paragraph(json["name"], json["type"])

        if "content" in json:
            if isinstance(json["content"], str):
                document.add_paragraph(json["content"], json["type"])
            elif isinstance(json["content"], list):
                for content in json["content"]:
                    Generator.generate_docx(document, content)
            else:
                Generator.generate_docx(document, json["content"])


d = {"general": {"date_start": "11/11/11",
                 "date_end": "12/12/12"},

     "clients": [{"name": "john doe",
                  "email": "john@doe.com",
                  "tel": "+33 66 55 44 33",
                  "role": "CEO"}],

     "auditors": [{"name": "haxor auditor",
                   "email": "haxor@auditor.com",
                   "tel": "31337",
                   "role": "pro-hacker"},

                  {"name": "script kiddie",
                   "email": "skiddie@mail.com",
                   "tel": "123456",
                   "role": "skiddie"}]}

p = (Generator.generate_json(d))
#print(p)
#print(Generator.generate_report(p))
doc = Document(docx="templates/template.docx")
Generator.generate_docx(doc, p)
doc.save("test.docx")